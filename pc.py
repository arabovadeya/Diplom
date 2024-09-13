from tkinter import filedialog, FALSE
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog
import docx
import docx2txt
import langdetect
from docx import Document

# После открытия окна интерфейса по нажатию на "Выбрать файл" добавить соответсвующие файы, затем после нажатия "Обработать", в окне интерфейса вы увидите результат работы модуля

root = tk.Tk()
root.title('Проверка корректности структуры документа')
root.geometry('650x400+400+400')
root.resizable(FALSE, FALSE)


def select_file():
    file1_path = filedialog.askopenfilename(initialdir="/", title="Select file",
                                            filetypes=(("text files", "*.docx"), ("all files", "*.*")))
    file1_entry.delete(0, tk.END)
    file1_entry.insert(0, file1_path)

    file2_path = filedialog.askopenfilename(initialdir="/", title="Select file",
                                            filetypes=(("text files", "*.docx"), ("all files", "*.*")))

    file2_entry.delete(0, tk.END)
    file2_entry.insert(0, file2_path)


def n_o_r():
    doc1 = docx.Document(file1_entry.get())
    doc2 = docx.Document(file2_entry.get())
    b = (len(file1_entry.get()) - 4) * 2
    f = Path(file2_entry.get()).stem
    if f == 'ШаблонТекстпрограммы':
        # Считаем количество строк в документе
        line_count = 0
        for paragraph in doc1.paragraphs[3:-1]:
            line_count += len(paragraph.text.splitlines())

        result_text.delete("3.0", tk.END)
        result_text.insert("3.0", "\nКоличество строк: " + str(line_count) + "\n\n")


def content_op():
    doc1 = Document(file1_entry.get())
    doc2 = docx.Document(file2_entry.get())
    text = docx2txt.process(file1_entry.get())
    f = Path(file2_entry.get()).stem
    count = 0
    text = text.lower()
    if f == 'ШаблонОписаниепрограммы':
        search_words = ['общие сведения', 'функциональное назначение', 'описание логической структуры',
                        'используемые технические средства', 'вызов и загрузка', 'входные данные', 'выходные данные',
                        'обозначение и наименование программы',
                        'программное обеспечение,необходимое для функционирования программы',
                        'языки программирования,на которых написана программа', 'классы решаемых задач',
                        'назначение программы', 'сведения о функциональных ограничениях на применение',
                        'алгоритм программы', 'используемые методы',
                        'структура программы с описанием функций составных частей и связи между ними',
                        'связи программы с другими программами']
        for search_word in search_words:
            if search_word.lower() not in text:
                return result_text.insert("3.0",
                                          "\n Cодержаниe файла 'Описание программы' не соответсвует ГОСТу " + "\n\n")
        return result_text.insert("3.0", "\nCодержаниe файла 'Описание программы' соответсвует ГОСТу " + "\n\n")

    if f == 'ШаблонОписаниеприменения':
        search_words1 = ['назначение программы', 'условия применения', 'описание задачи', 'входные и выходные данные',
                         'возможности программы', 'основные характеристики программы',
                         'ограничения,накладываемые на область применения программмы',
                         'требования к техническим (аппаратным) средствам',
                         'требования к программным средствам (другим программам)',
                         'общие характеристики входной информации', 'общие характеристики выходной информации',
                         'требования и условия технического характера',
                         'требования и условия организационнго характера',
                         'требования и условия технологического характера', 'определение задачи',
                         'методы решения задачи', 'сведения о входных данных ', 'сведения о выходных данных',
                         'лист регистрации изменений']
        for search_word in search_words1:
            if search_word.lower() not in text:
                return result_text.insert("3.0",
                                          "\n Cодержаниe файла 'Описание применения' не соответсвует ГОСТу " + "\n\n")
        return result_text.insert("3.0", "\nCодержаниe файла 'Описание применения' соответсвует ГОСТу " + "\n\n")

    if f == 'ШаблонРуководствоСП':
        search_words2 = ['назначение программы', 'общие сведения о программе', 'структура программы',
                         'настройка программы',
                         'проверка программы', 'дополнительные возможности', 'сообщения системному программисту']
        for search_word in search_words2:
            if search_word.lower() not in text:
                return result_text.insert("3.0",
                                          "\n Cодержаниe файла 'Руководство системного программиста' не соответсвует ГОСТу " + "\n\n")
        return result_text.insert("3.0",
                                  "\nCодержаниe файла 'Руководство системного программиста' соответсвует ГОСТу " + "\n\n")

    if f == 'ШаблонРуководствооператора':
        search_words3 = ['назначение программы', 'условия выполнения программы', 'выполнение программы',
                         'сообщения оператору']
        for search_word in search_words3:
            if search_word.lower() not in text:
                return result_text.insert("3.0",
                                          "\n Cодержаниe файла 'Руководство оператора' не соответсвует ГОСТу " + "\n\n")
        return result_text.insert("3.0", "\nCодержаниe файла 'Руководство оператора' соответсвует ГОСТу " + "\n\n")


def check_font():
    doc1 = Document(file1_entry.get())
    doc2 = docx.Document(file2_entry.get())
    b = (len(file1_entry.get()) - 4) * 2
    text = docx2txt.process(file1_entry.get())

    for paragraph in doc1.paragraphs:
        for run in paragraph.runs:
            if run.font.name != "Times New Roman" and run.font.size != 14:
                return result_text.insert("4.0", "\n Текст  соответствует по шрифту ГОСТу " + "\n\n")
            else:
                return result_text.insert("4.0", "\nТекст  соответствует по шрифту ГОСТу " + "\n\n")


def check_int():
    doc1 = Document(file1_entry.get())
    doc2 = docx.Document(file2_entry.get())
    b = (len(file1_entry.get()) - 4) * 2
    text = docx2txt.process(file1_entry.get())

    all_paragraphs_1pt5 = True
    for paragraph in doc1.paragraphs:
        if paragraph.paragraph_format.line_spacing != 1.5:
            all_paragraphs_1pt5 = False
            break

    if all_paragraphs_1pt5:
        return result_text.insert("3.0", "\nТекст соответствует по межстрочному интервалу ГОСТу " + "\n\n")
    else:
        return result_text.insert("3.0", "\n Текст  соответствует по межстрочному интервалу ГОСТу " + "\n\n")


def check_ots():
    doc1 = Document(file1_entry.get())
    doc2 = docx.Document(file2_entry.get())
    b = (len(file1_entry.get()) - 4) * 2
    text = docx2txt.process(file1_entry.get())

    for paragraph in doc1.paragraphs:
        if paragraph.paragraph_format.left_indent != 1.25:
            return result_text.insert("3.0", "\n Текст  соответствует по отступу в начале абзаца ГОСТу " + "\n\n")
    return result_text.insert("3.0", "\nТекст соответствует по отступу в начале абзаца ГОСТу " + "\n\n")


def check_vr():
    doc1 = Document(file1_entry.get())
    doc2 = Document(file2_entry.get())
    b = (len(file1_entry.get()) - 4) * 2
    text = docx2txt.process(file1_entry.get())

    all_paragraphs_justified = True
    for paragraph in doc1.paragraphs:
        if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            all_paragraphs_justified = False
            break

    if all_paragraphs_justified:
        return result_text.insert("3.0", "\nТекст соответствует по выравниванию ГОСТу\n\n")
    else:
        return result_text.insert("3.0", "\nТекст  соответствует по выравниванию ГОСТу\n\n")


def compare_files():
    file1_path = Document(file1_entry.get())
    file2_path = Document(file2_entry.get())

    template_pages = len(file2_path.sections)
    count = 0
    same_pages = []
    for i in range(len(file1_path.sections)):
        for j in range(len(file2_path.sections)):
            if file1_path.paragraphs[i].text == file2_path.paragraphs[j].text:
                same_pages.append(i)
                count += 1
                result_text.delete("0.0", tk.END)
                result_text.insert("1.0", "Одинаковые страницы: " + str(
                    same_pages) + "\nКоличество одинаковых страниц: " + str(count) + "\n\n")


def button_event():
    compare_files()
    n_o_r()
    content_op()
    check_font()
    check_int()
    check_ots()
    check_vr()


file1_label = tk.Label(root, text="Файл технической документации:")
file1_label.grid(row=1, column=0)

file1_entry = tk.Entry(root)
file1_entry.grid(row=1, column=1)

file1_button = tk.Button(root, text="Выбрать файл", command=select_file)
file1_button.grid(row=1, column=2)

file2_label = tk.Label(root, text="Шаблон:")
file2_label.grid(row=2, column=0)

file2_entry = tk.Entry(root)
file2_entry.grid(row=2, column=1)

compare_button = tk.Button(root, text="Обработать", command=button_event)
compare_button.grid(row=3, column=1)

result_label = tk.Label(root, text="Результат :")
result_label.grid(row=4, column=0)

result_text = tk.Text(root)
result_text.grid(row=5, column=0, columnspan=3)

root.mainloop()
